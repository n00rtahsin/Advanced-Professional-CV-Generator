import json
import os
from tkinter import Tk, Label, Entry, Text, Button, END, filedialog, StringVar, messagebox, Scale, HORIZONTAL, Frame, Listbox, ttk

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class CVApp:
    def __init__(self, master):
        self.master = master
        master.title("Professional CV Generator")
        master.configure(bg="#f0f8ff")

        self.data = {
            "skills": [],
            "job_experiences": [],
            "education": []
        }

        self.create_widgets()

    def create_widgets(self):
        # Create a frame for better organization
        main_frame = Frame(self.master, bg="#f0f8ff")
        main_frame.pack(padx=10, pady=10)

        # Create labels and entries for personal information
        self.create_label("Full Name", 0, main_frame)
        self.name = StringVar()
        self.create_entry(self.name, 0, main_frame)

        self.create_label("Title", 1, main_frame)
        self.title = StringVar()
        self.create_entry(self.title, 1, main_frame)

        self.create_label("Email", 2, main_frame)
        self.email = StringVar()
        self.create_entry(self.email, 2, main_frame)

        self.create_label("Phone", 3, main_frame)
        self.phone = StringVar()
        self.create_entry(self.phone, 3, main_frame)

        self.create_label("Address", 4, main_frame)
        self.address = StringVar()
        self.create_entry(self.address, 4, main_frame)

        self.create_label("LinkedIn", 5, main_frame)
        self.linkedin = StringVar()
        self.create_entry(self.linkedin, 5, main_frame)

        self.create_label("GitHub", 6, main_frame)
        self.github = StringVar()
        self.create_entry(self.github, 6, main_frame)

        self.create_label("Website", 7, main_frame)
        self.website = StringVar()
        self.create_entry(self.website, 7, main_frame)

        self.create_label("Profile Picture", 8, main_frame)
        self.pic_path = StringVar()
        Entry(main_frame, textvariable=self.pic_path, width=38).grid(row=8, column=1, sticky='w', padx=10)
        Button(main_frame, text="Browse", command=self.browse_pic, bg="#2196F3", fg="white").grid(row=8, column=1, sticky='e', padx=10)

        self.create_label("Summary", 9, main_frame)
        self.summary_text = Text(main_frame, height=4, width=50, wrap='word')
        self.summary_text.grid(row=9, column=1, padx=10)

        # Skills Section
        self.create_label("Skills", 10, main_frame)
        self.skills_frame = Frame(main_frame, bg="#f0f8ff")
        self.skills_frame.grid(row=10, column=1, padx=10, sticky='w')

        self.skill_name = StringVar()
        self.create_entry(self.skill_name, 0, self.skills_frame)

        self.skill_level = Scale(self.skills_frame, from_=1, to=10, orient=HORIZONTAL, bg="#f0f8ff")
        self.skill_level.grid(row=0, column=1, padx=10)

        Button(self.skills_frame, text="Add Skill", command=self.add_skill, bg="#4CAF50", fg="white").grid(row=0, column=2, padx=10)

        self.skills_listbox = Listbox(main_frame, width=50, height=5)
        self.skills_listbox.grid(row=11, column=1, padx=10)

        # Job Experience Section
        self.create_label("Job Title", 12, main_frame)
        self.job_title = StringVar()
        self.create_entry(self.job_title, 12, main_frame)

        self.create_label("Company", 13, main_frame)
        self.company_name = StringVar()
        self.create_entry(self.company_name, 13, main_frame)

        self.create_label("Job Description", 14, main_frame)
        self.job_description = Text(main_frame, height=4, width=50, wrap='word')
        self.job_description.grid(row=14, column=1, padx=10)

        Button(main_frame, text="Add Job Experience", command=self.add_job_experience, bg="#4CAF50", fg="white").grid(row=15, column=1, pady=5)

        # Display Job Experiences
        self.create_label("Job Experiences", 16, main_frame)
        self.job_experience_listbox = Listbox(main_frame, width=50, height=5)
        self.job_experience_listbox.grid(row=17, column=1, padx=10)

        # Education Section
        self.create_label("Degree", 18, main_frame)
        self.degree_name = StringVar()
        self.create_entry(self.degree_name, 18, main_frame)

        self.create_label("Institution", 19, main_frame)
        self.institution_name = StringVar()
        self.create_entry(self.institution_name, 19, main_frame)

        self.create_label("Year", 20, main_frame)
        self.year_name = StringVar()
        self.create_entry(self.year_name, 20, main_frame)

        Button(main_frame, text="Add Education", command=self.add_education, bg="#4CAF50", fg="white").grid(row=21, column=1, pady=5)

        # Display Education
        self.create_label("Education", 22, main_frame)
        self.education_listbox = Listbox(main_frame, width=50, height=5)
        self.education_listbox.grid(row=23, column=1, padx=10)

        # Hobbies Section
        self.create_label("Hobbies (comma separated)", 24, main_frame)
        self.hobbies_entry = StringVar()
        self.create_entry(self.hobbies_entry, 24, main_frame)

        # Template Selection
        self.create_label("Select Template", 25, main_frame)
        self.template_var = StringVar(value="Template 1")
        self.template_options = ["Template 1", "Template 2", "Template 3"]
        self.template_menu = ttk.Combobox(main_frame, textvariable=self.template_var, values=self.template_options)
        self.template_menu.grid(row=25, column=1, padx=10)

        # Generate CV Button
        Button(main_frame, text="Generate CV", bg="#4CAF50", fg="white", command=self.generate_cv).grid(row=26, column=1, pady=15)

    def create_label(self, text, row, frame):
        Label(frame, text=text, bg="#f0f8ff", font=("Segoe UI", 10, "bold"), fg="#333").grid(row=row, column=0, sticky='w', padx=10, pady=2)

    def create_entry(self, var, row, frame):
        entry = Entry(frame, textvariable=var, width=50)
        entry.grid(row=row, column=1, padx=10)
        return entry

    def browse_pic(self):
        filename = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if filename:
            self.pic_path.set(filename)

    def add_skill(self):
        skill = self.skill_name.get()
        level = self.skill_level.get()
        if skill:
            self.data['skills'].append({'name': skill, 'level': level})
            self.skills_listbox.insert(END, f"{skill} - Level {level}")
            self.create_progress_bar(skill, level)  # Add progress bar for the skill
            self.skill_name.set("")  # Clear the skill entry
            self.skill_level.set(1)  # Reset skill level
        else:
            messagebox.showwarning("Input Error", "Please enter a skill.")

    def create_progress_bar(self, skill_name, level):
        progress_frame = Frame(self.skills_frame, bg="#f0f8ff")
        progress_frame.grid(row=self.skills_listbox.size(), column=0, sticky='w', padx=10)

        Label(progress_frame, text=skill_name, bg="#f0f8ff").pack(side='left')
        progress = ttk.Progressbar(progress_frame, length=200, maximum=10, value=level)
        progress.pack(side='left', padx=5)

    def add_job_experience(self):
        job_title = self.job_title.get()
        company = self.company_name.get()
        description = self.job_description.get("1.0", END).strip()
        if job_title and company and description:
            job_experience = {
                'title': job_title,
                'company': company,
                'description': description
            }
            self.data['job_experiences'].append(job_experience)
            self.job_experience_listbox.insert(END, f"{job_title} at {company}")
            messagebox.showinfo("Success", "Job experience added.")
            self.job_title.set("")  # Clear the job title
            self.company_name.set("")  # Clear the company name
            self.job_description.delete("1.0", END)  # Clear the job description
        else:
            messagebox.showwarning("Input Error", "Please fill in all job experience fields.")

    def add_education(self):
        degree = self.degree_name.get()
        institution = self.institution_name.get()
        year = self.year_name.get()
        if degree and institution and year:
            education_entry = {
                'degree': degree,
                'institution': institution,
                'year': year
            }
            self.data['education'].append(education_entry)
            self.education_listbox.insert(END, f"{degree} from {institution} ({year})")
            messagebox.showinfo("Success", "Education added.")
            self.degree_name.set("")  # Clear the degree entry
            self.institution_name.set("")  # Clear the institution entry
            self.year_name.set("")  # Clear the year entry
        else:
            messagebox.showwarning("Input Error", "Please fill in all education fields.")

    def generate_cv(self):
        self.data['name'] = self.name.get()
        self.data['title'] = self.title.get()
        self.data['email'] = self.email.get()
        self.data['phone'] = self.phone.get()
        self.data['address'] = self.address.get()
        self.data['linkedin'] = self.linkedin.get()
        self.data['github'] = self.github.get()
        self.data['website'] = self.website.get()
        self.data['profile_pic'] = self.pic_path.get() if os.path.exists(self.pic_path.get()) else None
        self.data['summary'] = self.summary_text.get("1.0", END).strip()
        self.data['hobbies'] = [hobby.strip() for hobby in self.hobbies_entry.get().split(',') if hobby.strip()]
        self.data['template'] = self.template_var.get()

        try:
            with open("cv_data.json", "w") as f:
                json.dump(self.data, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error Saving JSON", str(e))
            return

        self.create_docx(self.data)

    def create_docx(self, data, filename=None):
        try:
            if not filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{data['name'].replace(' ', '_')}_CV_{timestamp}.docx"

            doc = Document()

            if data['profile_pic']:
                try:
                    doc.add_picture(data['profile_pic'], width=Inches(1.25))
                except Exception as e:
                    messagebox.showwarning("Image Warning", f"Profile picture could not be added: {e}")

            doc.add_heading(data['name'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(data['title']).alignment = WD_ALIGN_PARAGRAPH.CENTER

            contact = f"Email: {data['email']} | Phone: {data['phone']}\nAddress: {data['address']}"
            if data['linkedin']: contact += f"\nLinkedIn: {data['linkedin']}"
            if data['github']: contact += f"\nGitHub: {data['github']}"
            if data['website']: contact += f"\nWebsite: {data['website']}"
            doc.add_paragraph(contact)

            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(data['summary'])

            doc.add_heading("Technical Skills", level=1)
            for skill in data['skills']:
                doc.add_paragraph(f"{skill['name']} - Level {skill['level']}")

            doc.add_heading("Job Experiences", level=1)
            for job in data['job_experiences']:
                doc.add_heading(job['title'], level=2)
                doc.add_paragraph(f"Company: {job['company']}")
                doc.add_paragraph(job['description'])

            doc.add_heading("Education", level=1)
            for edu in data['education']:
                doc.add_paragraph(f"{edu['degree']} from {edu['institution']} ({edu['year']})")

            doc.add_heading("Hobbies", level=1)
            doc.add_paragraph(", ".join(data['hobbies']))

            doc.save(filename)
            messagebox.showinfo("Success", f"CV saved as {filename}")
        except Exception as e:
            messagebox.showerror("Error Saving DOCX", str(e))

def main():
    root = Tk()
    app = CVApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
