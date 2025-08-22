# Advanced Professional Python CV Generator with Colorful GUI and Skill Slider (Micropip-free)
import json
import os
from tkinter import Tk, Label, Entry, Text, Button, END, filedialog, StringVar, ttk, messagebox, Scale, HORIZONTAL
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class CVApp:
    def __init__(self, master):
        self.master = master
        master.title("Professional CV Generator")
        master.configure(bg="#f0f8ff")

        self.data = {}

        def create_label(text, row):
            Label(master, text=text, bg="#f0f8ff", font=("Segoe UI", 10, "bold"), fg="#333").grid(row=row, column=0, sticky='w', padx=10, pady=2)

        def create_entry(var, row):
            entry = Entry(master, textvariable=var, width=50)
            entry.grid(row=row, column=1, padx=10)
            return entry

        self.name = StringVar()
        self.title = StringVar()
        self.email = StringVar()
        self.phone = StringVar()
        self.address = StringVar()
        self.linkedin = StringVar()
        self.github = StringVar()
        self.website = StringVar()
        self.pic_path = StringVar()

        create_label("Full Name", 0)
        create_entry(self.name, 0)

        create_label("Title", 1)
        create_entry(self.title, 1)

        create_label("Email", 2)
        create_entry(self.email, 2)

        create_label("Phone", 3)
        create_entry(self.phone, 3)

        create_label("Address", 4)
        create_entry(self.address, 4)

        create_label("LinkedIn", 5)
        create_entry(self.linkedin, 5)

        create_label("GitHub", 6)
        create_entry(self.github, 6)

        create_label("Website", 7)
        create_entry(self.website, 7)

        create_label("Profile Picture", 8)
        Entry(master, textvariable=self.pic_path, width=38).grid(row=8, column=1, sticky='w', padx=10)
        Button(master, text="Browse", command=self.browse_pic).grid(row=8, column=1, sticky='e', padx=10)

        create_label("Summary", 9)
        self.summary_text = Text(master, height=4, width=50, wrap='word')
        self.summary_text.grid(row=9, column=1, padx=10)

        create_label("Skills (comma separated)", 10)
        self.skills_entry = Entry(master, width=50)
        self.skills_entry.grid(row=10, column=1, padx=10)

        create_label("Skill Level (1-10)", 11)
        self.skill_level = Scale(master, from_=1, to=10, orient=HORIZONTAL)
        self.skill_level.grid(row=11, column=1, padx=10, sticky='w')

        create_label("Languages", 12)
        self.languages_entry = Entry(master, width=50)
        self.languages_entry.grid(row=12, column=1, padx=10)

        Button(master, text="Generate CV", bg="#4CAF50", fg="white", font=("Segoe UI", 10, "bold"), command=self.generate_cv).grid(row=13, column=1, pady=15)

    def browse_pic(self):
        filename = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if filename:
            self.pic_path.set(filename)

    def generate_cv(self):
        self.data['name'] = self.name.get()
        self.data['title'] = self.title.get()
        self.data['email'] = self.email.get()
        self.data['phone'] = self.phone.get()
        self.data['address'] = self.address.get()
        self.data['linkedin'] = self.linkedin.get()
        self.data['github'] = self.github.get()
        self.data['website'] = self.website.get()
        self.data['profile_pic'] = self.pic_path.get() if os.path.exists(self.pic_path.get()) else None
        self.data['summary'] = self.summary_text.get("1.0", END).strip()
        self.data['skills'] = [s.strip() for s in self.skills_entry.get().split(',') if s.strip()]
        self.data['skill_rating'] = self.skill_level.get()
        self.data['languages'] = [l.strip() for l in self.languages_entry.get().split(',') if l.strip()]

        try:
            with open("cv_data.json", "w") as f:
                json.dump(self.data, f, indent=4)
        except Exception as e:
            messagebox.showerror("Error Saving JSON", str(e))
            return

        self.create_docx(self.data)

    def create_docx(self, data, filename=None):
        try:
            if not filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{data['name'].replace(' ', '_')}_CV_{timestamp}.docx"

            doc = Document()

            if data['profile_pic']:
                try:
                    doc.add_picture(data['profile_pic'], width=Inches(1.25))
                except Exception as e:
                    messagebox.showwarning("Image Warning", f"Profile picture could not be added: {e}")

            doc.add_heading(data['name'], 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(data['title']).alignment = WD_ALIGN_PARAGRAPH.CENTER

            contact = f"Email: {data['email']} | Phone: {data['phone']}\nAddress: {data['address']}"
            if data['linkedin']: contact += f"\nLinkedIn: {data['linkedin']}"
            if data['github']: contact += f"\nGitHub: {data['github']}"
            if data['website']: contact += f"\nWebsite: {data['website']}"
            doc.add_paragraph(contact)

            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(data['summary'])

            doc.add_heading("Technical Skills", level=1)
            skills_text = ", ".join(data['skills']) + f"\nSkill Proficiency: {data['skill_rating']}/10"
            doc.add_paragraph(skills_text)

            doc.add_heading("Languages", level=1)
            doc.add_paragraph(", ".join(data['languages']))

            doc.save(filename)
            messagebox.showinfo("Success", f"CV saved as {filename}")
        except Exception as e:
            messagebox.showerror("Error Saving DOCX", str(e))


def main():
    root = Tk()
    app = CVApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
