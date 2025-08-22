import json
from docx import Document
from docx.shared import Pt
import os

# Optional PDF conversion
try:
    from docx2pdf import convert
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def collect_user_info():
    print("=== Welcome to Smart CV Builder ===\nPlease enter the following information:\n")

    data = {}

    # Personal Info
    data['name'] = input("Full Name: ")
    data['email'] = input("Email: ")
    data['phone'] = input("Phone Number: ")
    data['address'] = input("Address: ")
    data['linkedin'] = input("LinkedIn URL (optional): ")
    data['github'] = input("GitHub URL (optional): ")

    # Career Objective
    print("\nEnter your Career Objective (Press ENTER twice to finish):")
    lines = []
    while True:
        line = input()
        if line == "":
            break
        lines.append(line)
    data['objective'] = "\n".join(lines)

    # Education
    print("\n--- Educational Background ---")
    education = []
    while True:
        print("Add a new education entry:")
        degree = input("Degree (e.g., BSc in CSE): ")
        institute = input("Institution: ")
        year = input("Year of Graduation: ")
        result = input("CGPA/Result: ")
        education.append({
            "degree": degree,
            "institute": institute,
            "year": year,
            "result": result
        })
        more = input("Add another education? (y/n): ")
        if more.lower() != 'y':
            break
    data['education'] = education

    # Skills
    print("\nEnter your skills (comma separated):")
    skills = input("Skills: ").split(',')
    data['skills'] = [skill.strip() for skill in skills]

    # Experience
    print("\nDo you have work experience? (y/n)")
    if input().lower() == 'y':
        experience = []
        while True:
            print("Add a new job entry:")
            title = input("Job Title: ")
            company = input("Company: ")
            duration = input("Duration (e.g., Jan 2020 - Dec 2022): ")
            print("Responsibilities (Press ENTER twice to finish):")
            res_lines = []
            while True:
                res = input()
                if res == "":
                    break
                res_lines.append(res)
            experience.append({
                "title": title,
                "company": company,
                "duration": duration,
                "responsibilities": res_lines
            })
            more = input("Add another job? (y/n): ")
            if more.lower() != 'y':
                break
        data['experience'] = experience
    else:
        data['experience'] = []

    # Save as JSON
    with open("cv_data.json", "w") as f:
        json.dump(data, f, indent=4)

    return data


def generate_docx(data, filename="final_cv.docx"):
    doc = Document()

    # Name as Heading
    doc.add_heading(data['name'], level=0)

    # Contact Info
    contact = f"Email: {data['email']} | Phone: {data['phone']}\nAddress: {data['address']}"
    if data['linkedin']:
        contact += f"\nLinkedIn: {data['linkedin']}"
    if data['github']:
        contact += f"\nGitHub: {data['github']}"
    doc.add_paragraph(contact)

    # Objective
    doc.add_heading("Career Objective", level=1)
    doc.add_paragraph(data['objective'])

    # Education
    doc.add_heading("Education", level=1)
    for edu in data['education']:
        entry = f"{edu['degree']} - {edu['institute']} ({edu['year']})\nResult: {edu['result']}"
        doc.add_paragraph(entry)

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data['skills']))

    # Experience
    if data['experience']:
        doc.add_heading("Work Experience", level=1)
        for exp in data['experience']:
            doc.add_paragraph(f"{exp['title']} - {exp['company']} ({exp['duration']})", style='List Bullet')
            for res in exp['responsibilities']:
                doc.add_paragraph(f"- {res}", style='List Bullet 2')

    doc.save(filename)
    print(f"\n✅ CV saved as {filename}")


def convert_to_pdf(docx_file="final_cv.docx", pdf_file="final_cv.pdf"):
    if not PDF_AVAILABLE:
        print("⚠️ PDF generation skipped (docx2pdf not installed).")
        return
    try:
        convert(docx_file, pdf_file)
        print(f"📄 PDF also saved as {pdf_file}")
    except Exception as e:
        print(f"❌ PDF conversion failed: {e}")


if __name__ == "__main__":
    user_data = collect_user_info()
    generate_docx(user_data)
    convert_to_pdf()
    print("\n🎉 CV Generation Completed!")
